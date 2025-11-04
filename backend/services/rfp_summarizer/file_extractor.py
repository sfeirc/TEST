"""
Service d'extraction de texte depuis des fichiers
Supporte: PDF, DOCX, TXT et autres formats de documents courants
"""
import os
from typing import Optional
import PyPDF2
from docx import Document
import chardet

def extract_text_from_pdf(file_path: str) -> str:
    """Extraire le texte d'un fichier PDF"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du texte PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extraire le texte d'un fichier DOCX"""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extraire également le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")

def extract_text_from_txt(file_path: str) -> str:
    """Extraire le texte d'un fichier TXT avec détection d'encodage"""
    try:
        # Détecter l'encodage
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'
        
        # Lire avec l'encodage détecté
        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
            return file.read().strip()
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du texte TXT: {str(e)}")

def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    Extraire le texte d'un fichier uploadé selon son extension
    
    Args:
        file_path: Chemin vers le fichier temporaire
        filename: Nom de fichier original avec extension
    
    Returns:
        Contenu textuel extrait
    """
    # Obtenir l'extension du fichier
    _, ext = os.path.splitext(filename.lower())
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.txt', '.text', '.md', '.markdown']:
        return extract_text_from_txt(file_path)
    else:
        # Essayer de lire comme texte par défaut
        try:
            return extract_text_from_txt(file_path)
        except:
            raise Exception(f"Format de fichier non supporté: {ext}. Formats supportés: PDF, DOCX, TXT")
